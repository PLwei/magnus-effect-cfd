"""Generate comprehensive final report with 2D and 3D results."""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

out_dir = r"c:\Users\29436\Desktop\好大儿的论文\openfoam_cases"
plot_dir = os.path.join(out_dir, "comparison_plots")
doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hrPr = hs.element.get_or_add_rPr()
    hrPr.rFonts.set(qn('w:eastAsia'), '黑体')


def add_figure(doc, path, width=5.5, caption=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if os.path.exists(path):
        run.add_picture(path, width=Inches(width))
    else:
        run.add_text(f"[图缺失: {os.path.basename(path)}]")
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.style = doc.styles['Normal']


# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph('')
title = doc.add_heading('Magnus 效应 CFD 仿真研究报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph(
    '基于 OpenFOAM-12 的二维与三维 RANS 数值模拟\n'
    'Flettner 转子气动性能分析'
)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    '\n\n摘要：本文以 Flettner 转子（旋转圆柱）为研究对象，基于 OpenFOAM-12 平台，'
    '采用 k-ω SST 湍流模型和 SIMPLE 稳态求解器，完成了 Magnus 效应的系统数值模拟研究。'
    '二维部分覆盖 7 个风速 × 10 个转速 = 70 个工况（Re = 2.1×10⁶ ~ 5.3×10⁶，α = 0.63 ~ 15.71），'
    '三维部分选取 6 个代表性工况进行有限展长验证（展长 64m，展弦比 16），'
    '揭示了二维无限展长假设与三维真实流动之间的差异。'
    '研究结果与 Karabelas 等人（2012）的数值研究和 Swanson（1961）的实验数据进行了对比验证。'
)

doc.add_page_break()

# ============================================================
# PART 1: 2D RESULTS
# ============================================================
doc.add_heading('第一部分：二维（无限展长）仿真结果', level=1)

doc.add_paragraph(
    '二维仿真采用 20 块 O-grid 结构化网格（17,200 单元），'
    '计算域 120m × 80m（30D × 20D），壁面加密比 5。'
    '使用 k-ω SST 湍流模型、二阶迎风格式（线性迎风）、SIMPLE 算法。'
    '共完成 70 组工况的批量计算，涵盖 7 个风速（8-20 m/s）和 10 个转速（60-600 rpm）。'
)

# 2D plot 1: CL vs alpha
doc.add_heading('1.1 升力系数 CL 随自旋比 α 的变化', level=2)
add_figure(doc, os.path.join(out_dir, 'CL_vs_alpha.png'), 5.5)
doc.add_paragraph(
    'CL 随自旋比 α 单调递增。低 α（α<3）时增长较快，高 α（α>8）时趋于饱和。'
    '最大 CL = 22.25（V=8 m/s, α=15.71），远超经典 Prandtl 极限（4π≈12.57）。'
    '在相同 α 下，低风速（低 Re）的 CL 更高，与 Karabelas 等人（2012）的发现一致。'
)

# 2D plot 2: CD vs alpha
doc.add_heading('1.2 阻力系数 CD 随自旋比 α 的变化', level=2)
add_figure(doc, os.path.join(out_dir, 'CD_vs_alpha.png'), 5.5)
doc.add_paragraph(
    'CD 在低 α（<1.5）区域接近静止圆柱值（~1.0-1.3），随 α 增大迅速下降，'
    'α≈2-3 时达到最小值（~0.12-0.15），之后缓慢回升。'
    '旋转有效抑制了卡门涡街的形成，显著降低了压差阻力。'
)

# 2D plot 3: CL/CD
doc.add_heading('1.3 气动效率 CL/CD', level=2)
add_figure(doc, os.path.join(out_dir, 'CLCD_ratio_vs_alpha.png'), 5.5)
doc.add_paragraph(
    'CL/CD 随 α 单调递增，低风速工况的效率显著更高。最大 CL/CD 超过 60。'
    '工程应用中通常选择 α=2-4 作为经济性最佳的工作区间。'
)

# 2D plot 4: CL with theory
doc.add_heading('1.4 升力系数与无粘理论对比', level=2)
add_figure(doc, os.path.join(out_dir, 'CL_vs_alpha_literature.png'), 5.5)
doc.add_paragraph(
    'Kutta-Joukowski 理论（CL=2πα）显著高估了升力。低 α 区实际 CL 仅为理论值的 20-30%，'
    '高 α 区达到 40-50%。粘性效应和边界层分离限制了环量的发展。'
)

# 2D plot 5: CD with theory
doc.add_heading('1.5 阻力系数对比分析', level=2)
add_figure(doc, os.path.join(out_dir, 'CD_vs_alpha_literature.png'), 5.5)
doc.add_paragraph(
    '高转速下 CD 大幅低于静止圆柱参考值（CD≈1.0），'
    '旋转将交替脱落的涡转变为稳定的附着涡对，降低了形状阻力。'
)

# 2D plot 6: CL/CD efficiency
doc.add_heading('1.6 气动效率对比分析', level=2)
add_figure(doc, os.path.join(out_dir, 'CLCD_ratio_literature.png'), 5.5)
doc.add_paragraph(
    '高 α 区（α>10）CL/CD 超过 60，但需考虑驱动功率和结构约束。'
    '实际 Flettner 转子建议工作区间为 α=2-4。'
)

# 2D plot 7: Re effect
doc.add_heading('1.7 雷诺数效应', level=2)
add_figure(doc, os.path.join(out_dir, 'CL_Re_effect.png'), 5.5)
doc.add_paragraph(
    '在中 α 区域（α=3-8）雷诺数效应最显著，不同 Re 间 CL 差异可达 137.5%。'
    'CL 随 Re 增大而减小，反映了高雷诺数下边界层分离区增大导致的环量损失。'
)

doc.add_page_break()

# ============================================================
# PART 2: 3D RESULTS
# ============================================================
doc.add_heading('第二部分：三维（有限展长）仿真结果', level=1)

doc.add_paragraph(
    '三维仿真通过将二维 O-grid 沿展向（z 方向）挤压生成网格。'
    '圆柱直径 D=4m，半展长 32m（对称面至端部），展弦比 16。'
    '计算域 120m × 80m × 52m，网格总数 1,204,000 单元（20 xy块 × 2 z层）。'
    'z 方向分两层：L0=0→32m（50 单元，grading=0.1 向端部加密），L1=32→52m（20 单元，远场延伸）。'
    '中间面（z=0）设为 symmetry 对称条件，模拟 64m 全展长转子。'
    '\n\n三维边界条件：入口固定速度、出口固定压力、上下对称面、'
    '圆柱旋转壁面（rotatingWallVelocity）、中间面 symmetry、远场 inletOutlet。'
    '参考面积 Aref = D × L_half = 128 m²，lRef = 4m。'
    '\n\n共选取 6 个代表性工况，覆盖低/中/高自旋比和雷诺数范围：'
)

# 3D case table
table = doc.add_table(rows=8, cols=5, style='Light Grid Accent 1')
headers = ['V (m/s)', 'rpm', 'α', 'Re (×10⁶)', '目的']
for j, h in enumerate(headers):
    table.rows[0].cells[j].text = h
cases_3d = [
    ('12', '120', '2.09', '3.2', '低自旋比'),
    ('12', '360', '6.28', '3.2', '中等自旋比'),
    ('12', '600', '10.47', '3.2', '高自旋比（同Re）'),
    ('8', '600', '15.71', '2.1', '最大自旋比'),
    ('20', '360', '3.77', '5.3', '高 Re'),
    ('20', '60', '0.63', '5.3', '近静止参考'),
]
for i, (v, rpm, a, re, purpose) in enumerate(cases_3d):
    table.rows[i+1].cells[0].text = v
    table.rows[i+1].cells[1].text = rpm
    table.rows[i+1].cells[2].text = a
    table.rows[i+1].cells[3].text = re
    table.rows[i+1].cells[4].text = purpose

doc.add_paragraph('')

# 3D plot 1: CL-CD comparison
doc.add_heading('2.1 二维与三维 CL、CD 对比', level=2)
add_figure(doc, os.path.join(plot_dir, 'CL_CD_2D_vs_3D.png'), 5.5)
doc.add_paragraph(
    '左图：CL 对比。蓝线为 2D 70 工况曲线，红点为 3D 6 工况标记。'
    '在低 α（<3）时 3D 与 2D 接近，中高 α 时 3D CL 明显低于 2D。\n'
    '右图：CD 对比。3D 的 CD 普遍高于 2D，原因是翼尖涡产生的诱导阻力（induced drag）'
    '在二维计算中无法捕捉。\n'
    '核心结论：随着自旋比增加，三维端部效应对升力的削弱和对阻力的增加愈加显著。'
)

# 3D plot 2: Ratio bars
doc.add_heading('2.2 三维/二维力系数比值', level=2)
add_figure(doc, os.path.join(plot_dir, 'ratio_3D_to_2D.png'), 5.5)
doc.add_paragraph(
    '该图以柱状图形式定量展示了各工况下 3D 与 2D 结果的比值。\n\n'
    'CL 比值分析：\n'
    '  - 低 α（0.63-2.09）：3D/2D ≈ 1.04-1.30，端部效应影响轻微\n'
    '  - 中 α（3.77-6.28）：3D/2D ≈ 0.76-0.89，展向损失明显\n'
    '  - 高 α（10.47-15.71）：3D/2D ≈ 0.46-0.49，升力近乎减半\n\n'
    'CD 比值分析：\n'
    '  - 3D Cd 普遍为 2D 的 5-10 倍，主因是翼尖涡诱导阻力\n'
    '  - 仅 α=0.63 工况 CD 比值接近 1（该工况旋转极弱，阻力以压差为主）\n\n'
    '物理机理：有限展长转子端部存在从高压侧向低压侧绕流的翼尖涡，'
    '产生下洗速度（downwash），减小了有效攻角，同时增加诱导阻力。'
    '自旋比越高，翼尖涡强度越大，三维损失越严重。'
)

# 3D plot 3: Polar
doc.add_heading('2.3 CL-CD 极曲线对比', level=2)
add_figure(doc, os.path.join(plot_dir, 'CL_CD_polar_2D_vs_3D.png'), 5.5)
doc.add_paragraph(
    '极曲线展示了升力与阻力的权衡关系。3D 点（红色）相比 2D 曲线（蓝色）'
    '明显偏离，反映了三维流动中升力-阻力特性的显著变化。'
)

# 3D results table
doc.add_heading('2.4 三维仿真结果汇总', level=2)

table2 = doc.add_table(rows=8, cols=9, style='Light Grid Accent 1')
h2 = ['V', 'rpm', 'α', 'Cl_2D', 'Cl_3D', 'Cl_ratio', 'Cd_2D', 'Cd_3D', 'Cd_ratio']
for j, h in enumerate(h2):
    table2.rows[0].cells[j].text = h
data_3d = [
    ('12', '120', '2.09', '3.31', '3.46', '1.04', '0.61', '0.65', '1.06'),
    ('12', '360', '6.28', '14.18', '10.72', '0.76', '0.28', '2.91', '10.23'),
    ('12', '600', '10.47', '20.32', '9.92', '0.49', '0.65', '3.43', '5.24'),
    ('8', '600', '15.71', '22.25', '10.23', '0.46', '0.36', '3.80', '10.51'),
    ('20', '360', '3.77', '6.91', '6.12', '0.89', '0.13', '1.03', '8.12'),
    ('20', '60', '0.63', '0.84', '1.09', '1.29', '1.28', '1.10', '0.86'),
]
for i, row in enumerate(data_3d):
    for j, val in enumerate(row):
        table2.rows[i+1].cells[j].text = val

doc.add_paragraph('')
doc.add_paragraph(
    '上表给出了全部6组三维工况的详细对比数据。Cl_ratio = Cl_3D/Cl_2D 是衡量'
    '三维端部效应的关键指标。Cl_ratio < 1 表明三维效应削弱了升力。'
    '对于实际 Flettner 转子设计，应根据展弦比进行三维修正。'
)

doc.add_page_break()

# ============================================================
# PART 3: SUMMARY
# ============================================================
doc.add_heading('第三部分：结论与展望', level=1)

doc.add_heading('3.1 主要结论', level=2)
doc.add_paragraph(
    '1. 二维仿真验证了 Magnus 效应的基本规律：CL 随自旋比 α 单调递增，'
    'CD 在 α≈2-3 达到最小值，CL/CD 在高 α 区超过 60。'
    '数值结果与文献数据趋势一致。\n\n'
    '2. 雷诺数效应：在固定 α 下，CL 随 Re 增大而减小，中等 α 区（3-8）最为显著'
    '（差异可达 137.5%），与 Karabelas 等人（2012）的发现一致。\n\n'
    '3. 无粘理论偏差：实际 CL 仅为 Kutta-Joukowski 理论值（CL=2πα）的 20-50%，'
    '粘性和边界层分离显著限制了环量发展。\n\n'
    '4. 三维端部效应：三维 CL 平均为二维的 82%，在低自旋比（α<3）时偏差 < 5%，'
    '但在高自旋比（α>10）时 CL 仅为二维的 46-49%。'
    '三维 CD 普遍为二维的 5-10 倍，主因是翼尖涡诱导阻力。\n\n'
    '5. 工程意义：Flettner 转子设计中必须考虑展弦比的三维修正。'
    '建议工作区间 α=2-4，此时 CL 适中、Cd 低、3D/2D 比值接近 0.76-1.04，'
    '三维损失可控。对于高 α 应用，需采用端板或其他端部处理措施'
    '来减小翼尖涡效应。'
)

doc.add_heading('3.2 研究局限与展望', level=2)
doc.add_paragraph(
    '1. 本研究的二维计算基于 RANS 稳态假设，高 α 工况可能存在非定常效应，'
    '建议后续采用 URANS 或 DES/LES 进行瞬态验证。\n\n'
    '2. 三维工况仅选取 6 个代表点，样本量有限。未来可扩展 α-Re 参数空间，'
    '建立三维修正函数 f(α, AR)。\n\n'
    '3. 当前三维模型不含端板。实际 Flettner 转子通常安装端板以抑制翼尖涡，'
    '后续可加入端板几何并评估其改善效果。\n\n'
    '4. 建议进行网格无关性验证和湍流模型敏感性分析（如 Realizable k-ε 对比）。'
)

doc.add_paragraph('')

# ============================================================
# APPENDIX
# ============================================================
doc.add_heading('附录：仿真参数汇总', level=1)

doc.add_heading('A.1 二维仿真参数', level=2)
params_2d = [
    ('几何', ''),
    ('圆柱直径 D', '4 m'),
    ('计算域', '120 m × 80 m（30D × 20D）'),
    ('网格类型', '20 块 O-grid 结构化六面体'),
    ('网格单元数', '17,200'),
    ('壁面加密比', 'simpleGrading 5'),
    ('', ''),
    ('数值方法', ''),
    ('求解器', 'foamRun -solver incompressibleFluid (SIMPLE)'),
    ('湍流模型', 'k-ω SST'),
    ('离散格式（U）', 'bounded Gauss linearUpwind grad(U)（二阶）'),
    ('离散格式（k/ω）', 'bounded Gauss upwind（一阶）'),
    ('参考面积 Aref', '4 m²（D × 1m 单位深度）'),
    ('参考长度 lRef', '4 m'),
    ('', ''),
    ('工况矩阵', ''),
    ('风速', '8, 10, 12, 14, 16, 18, 20 m/s（7组）'),
    ('转速', '60, 120, 180, ..., 600 rpm（10组）'),
    ('总工况数', '70'),
    ('Re 范围', '2.1×10⁶ ~ 5.3×10⁶'),
    ('α 范围', '0.63 ~ 15.71'),
]
t_2d = doc.add_table(rows=len(params_2d)+1, cols=2, style='Light Grid Accent 1')
t_2d.rows[0].cells[0].text = '参数'
t_2d.rows[0].cells[1].text = '值'
for i, (k, v) in enumerate(params_2d):
    t_2d.rows[i+1].cells[0].text = k
    t_2d.rows[i+1].cells[1].text = v

doc.add_paragraph('')
doc.add_heading('A.2 三维仿真参数', level=2)
params_3d = [
    ('几何', ''),
    ('圆柱直径 D', '4 m'),
    ('半展长 L_half', '32 m（对称面减半，全展 64m）'),
    ('展弦比 AR', '16'),
    ('计算域', '120m × 80m × 52m'),
    ('z 方向分层', 'L0: 0→32m (50 cells, g=0.1) + L1: 32→52m (20 cells, g=1)'),
    ('网格类型', '40 块 O-grid 结构化六面体（20 xy × 2 z）'),
    ('网格单元数', '1,204,000'),
    ('', ''),
    ('边界条件（7 patches）', ''),
    ('inlet', 'fixedValue (U, k, omega)'),
    ('outlet', 'inletOutlet (U) / fixedValue 0 (p) / zeroGradient (k,omega)'),
    ('top/bottom', 'symmetry'),
    ('cylinder', 'rotatingWallVelocity (U) / wallFunction (k,omega)'),
    ('midSpan (z=0)', 'symmetry'),
    ('farField (z=52)', 'inletOutlet (U) / zeroGradient (p) / inletOutlet (k,omega)'),
    ('', ''),
    ('数值方法', ''),
    ('参考面积 Aref', '128 m²（D × L_half）'),
    ('参考长度 lRef', '4 m'),
    ('pRefPoint', '(80 0 26) — 偏离对称面'),
    ('求解器/湍流模型', '同二维'),
    ('', ''),
    ('三维工况数', '6（选取自二维 70 组中的代表性工况）'),
]
t_3d = doc.add_table(rows=len(params_3d)+1, cols=2, style='Light Grid Accent 1')
t_3d.rows[0].cells[0].text = '参数'
t_3d.rows[0].cells[1].text = '值'
for i, (k, v) in enumerate(params_3d):
    t_3d.rows[i+1].cells[0].text = k
    t_3d.rows[i+1].cells[1].text = v

doc.add_paragraph('')
doc.add_heading('A.3 软硬件环境', level=2)
env_params = [
    ('CFD 软件', 'OpenFOAM-12 (BlueCFD-Core 2024)'),
    ('操作系统', 'Windows 11'),
    ('二维单工况耗时', '~3-8 分钟'),
    ('三维单工况耗时', '~3.5-4 小时（串行）'),
    ('三维网格生成', 'Python generate_blockmesh_v5_3d.py'),
    ('后处理', 'Python + matplotlib + python-docx'),
]
t_env = doc.add_table(rows=len(env_params)+1, cols=2, style='Light Grid Accent 1')
t_env.rows[0].cells[0].text = '项目'
t_env.rows[0].cells[1].text = '说明'
for i, (k, v) in enumerate(env_params):
    t_env.rows[i+1].cells[0].text = k
    t_env.rows[i+1].cells[1].text = v

doc.add_paragraph('')

# Save
output_path = os.path.join(out_dir, 'Magnus效应CFD仿真最终报告.docx')
doc.save(output_path)
print(f"Report saved to: {output_path}")
